import io
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


def _add_horizontal_rule(doc: Document) -> None:
    """단락 아래에 가는 수평선을 그립니다 (단락 bottom border 방식)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "DDDDDD")
    pBdr.append(bottom)
    pPr.append(pBdr)


def export(summary: str, source_filename: str) -> bytes:
    doc = Document()

    # ── 제목 ──────────────────────────────────────────────────────────────────
    doc.add_heading("요약 결과", level=1)

    # ── 메타데이터 ─────────────────────────────────────────────────────────────
    meta_parts: list[str] = []
    if source_filename:
        meta_parts.append(f"원본 파일: {source_filename}")
    meta_parts.append(f"생성 일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    meta_p = doc.add_paragraph("  |  ".join(meta_parts))
    for run in meta_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    _add_horizontal_rule(doc)

    # ── 본문 ──────────────────────────────────────────────────────────────────
    # 줄바꿈 단위로 분리해 단락 구조를 유지합니다.
    for line in summary.split("\n"):
        p = doc.add_paragraph(line if line.strip() else "")
        p.style.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
