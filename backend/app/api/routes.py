# API 라우터 모듈
#
# 이 파일의 역할:
# 1. 입력 검증: 서비스 계층에 도달하기 전에 분명히 잘못된 요청을 걸러냅니다.
# 2. 에러 변환: 서비스에서 올라온 RuntimeError를 HTTP 상태 코드로 변환합니다.
#
# 비즈니스 로직은 summarizer.py에 작성되어 있습니다.

import io
import time
from datetime import datetime
from pathlib import Path
from urllib.parse import quote

from docx import Document
from docx.oxml.ns import qn
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import Response
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.dependencies import get_current_user
from app.db.models.user import User
from app.db.session import get_db
from app.repositories import summary_repository, download_repository
from app.schemas.summarize import SummarizeRequest, SummarizeResponse
from app.schemas.export import ExportRequest
from app.schemas.chat import ChatRequest, ChatResponse
from app.services import summarizer, export_service, chat_service
from app.services.ocr_extractor import extract_text_from_image
from app.services.pdf_extractor import extract_text_from_pdf

router = APIRouter()


def _build_export_filename(source_filename: str, ext: str) -> str:
    """
    다운로드 파일명을 생성합니다.

    source_filename이 있으면:  "{base}_요약결과_{date}.{ext}"
    source_filename이 없으면:  "요약결과_{date}.{ext}"

    base는 원본 파일명의 마지막 확장자만 제거합니다.
    예: "report.final.v2.docx" → "report.final.v2"
    """
    date_str = datetime.now().strftime("%Y%m%d")
    if source_filename:
        base = Path(source_filename).stem  # 마지막 확장자 하나만 제거
        return f"{base}_요약결과_{date_str}.{ext}"
    return f"요약결과_{date_str}.{ext}"

# Swagger UI나 테스트 도구의 기본값("string", "example" 등)처럼
# 의미 없는 입력을 차단합니다.
PLACEHOLDER_VALUES = {"string", "text", "example", "sample"}

# 파일 업로드에서 허용하는 확장자 목록입니다.
# 새 형식을 지원할 때 여기에 추가하고, _extract_text()에 분기를 추가합니다.
ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx", ".png", ".jpg", ".jpeg"}


# ── 파일별 텍스트 추출 헬퍼 ──────────────────────────────────────────────────
#
# 각 함수는 파일 형식에 맞게 텍스트를 추출하고, 실패 시 HTTPException을 raise합니다.
# 라우트 함수는 확장자에 따라 적절한 헬퍼를 호출

def _read_txt(raw: bytes) -> str:
    """
    txt 파일 바이트를 UTF-8 문자열로 디코딩합니다.
    utf-8-sig를 사용해 Windows 메모장 등이 추가하는 BOM도 자동 제거합니다.
    """
    try:
        return raw.decode("utf-8-sig").strip()
    except UnicodeDecodeError:
        # UTF-8이 아닌 인코딩(EUC-KR 등)으로 저장된 파일은 여기서 걸립니다.
        raise HTTPException(
            status_code=400,
            detail="파일을 UTF-8로 읽을 수 없습니다. 파일을 UTF-8로 저장한 뒤 다시 시도해 주세요.",
        )



def _read_docx(raw: bytes) -> str:
    """
    Word(.docx) 파일 바이트에서 텍스트를 추출합니다.

    지원 범위:
    - 일반 문단 및 제목 (Heading 스타일 포함 — 내부적으로 모두 w:p 요소)
    - 표(table) 셀 텍스트 — 행 단위로 추출, 셀은 " | "로 구분

    제외 범위 (1차 버전):
    - 머리글/바닥글 (페이지 번호·문서 제목 등 요약 불필요 정보)
    - 텍스트 상자·도형 안 텍스트 (추출 복잡도 대비 빈도 낮음)
    - 주석(comment), 각주(footnote), 이미지 내 텍스트

    문서 구조 보존:
    - doc.element.body를 직접 순회해 단락과 표를 원문 순서대로 처리합니다.
    - doc.paragraphs + doc.tables를 따로 수집하면 문서 내 순서가 뒤섞입니다.

    병합 셀 처리:
    - 가로 병합 셀은 python-docx 내부에서 동일 텍스트가 반복됩니다.
    - dict.fromkeys()로 삽입 순서를 유지하면서 중복을 제거합니다.
    """
    try:
        doc = Document(io.BytesIO(raw))
    except Exception:
        # python-docx는 zipfile.BadZipFile, ValueError 등 다양한 예외를 낼 수 있습니다.
        # .doc(구형 바이너리 포맷)을 올리면 여기서 걸립니다.
        raise HTTPException(
            status_code=400,
            detail=(
                "Word 파일을 읽을 수 없습니다. "
                "파일이 손상되었거나 올바른 .docx 형식이 아닙니다. "
                "구형 .doc 파일은 지원하지 않으며, .docx로 변환 후 업로드해 주세요."
            ),
        )

    parts: list[str] = []

    for block in doc.element.body:
        # namespace URI를 제거해 로컬 태그 이름만 비교합니다.
        # 예: "{http://...wordml...}p" → "p"
        tag = block.tag.split("}")[-1]

        if tag == "p":
            # 단락 및 제목: w:t 요소의 텍스트를 이어붙입니다.
            # run 분리나 스타일 구분 없이 순수 텍스트만 추출합니다.
            text = "".join(
                node.text or ""
                for node in block.iter()
                if node.tag.endswith("}t")
            ).strip()
            if text:
                parts.append(text)

        elif tag == "tbl":
            # 표: 행(w:tr) → 셀(w:tc) 순서로 순회합니다.
            for row in block.findall(".//" + qn("w:tr")):
                seen: dict[str, None] = {}  # 삽입 순서 유지 + 중복 제거
                for cell in row.findall(".//" + qn("w:tc")):
                    cell_text = "".join(
                        node.text or ""
                        for node in cell.iter()
                        if node.tag.endswith("}t")
                    ).strip()
                    if cell_text:
                        seen[cell_text] = None  # dict는 삽입 순서를 보장합니다 (Python 3.7+)
                if seen:
                    parts.append(" | ".join(seen.keys()))

    text = "\n\n".join(parts).strip()

    if not text:
        raise HTTPException(
            status_code=400,
            detail=(
                "Word 파일에서 텍스트를 추출할 수 없습니다. "
                "문서 내용이 비어 있거나, 텍스트 상자·이미지로만 구성된 파일일 수 있습니다."
            ),
        )

    return text


# ── 라우트 ────────────────────────────────────────────────────────────────────

@router.get("/health")
def health_check():
    return {"status": "ok"}


@router.post("/summarize", response_model=SummarizeResponse)
def summarize_route(
    request: SummarizeRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    text = request.text.strip()

    # ── 입력 검증 ────────────────────────────────────────────────────────────
    if not text:
        raise HTTPException(status_code=400, detail="텍스트가 비어 있습니다.")

    if text.lower() in PLACEHOLDER_VALUES:
        raise HTTPException(status_code=400, detail="실제 문서 내용을 입력해 주세요. 예시 텍스트는 허용되지 않습니다.")

    if len(text) < 10:
        raise HTTPException(status_code=400, detail="텍스트가 너무 짧습니다. 요약할 내용을 충분히 입력해 주세요.")

    # ── 요약 실행 ─────────────────────────────────────────────────────────────
    summary_mode = "chunked" if len(text) > settings.chunk_threshold else "single"
    start_ms = time.monotonic()

    try:
        result = summarizer.summarize(text)
    except RuntimeError as e:
        summary_repository.create(
            db,
            user_id=current_user.user_id,
            model_name=settings.ollama_model,
            summary_mode=summary_mode,
            input_chars=len(text),
            output_summary="",
            status="FAILED",
            file_type="text",
            document_type=request.document_type,
            error_message=str(e),
            processing_time_ms=int((time.monotonic() - start_ms) * 1000),
        )
        raise HTTPException(status_code=502, detail=str(e))

    record = summary_repository.create(
        db,
        user_id=current_user.user_id,
        model_name=settings.ollama_model,
        summary_mode=summary_mode,
        input_chars=len(text),
        output_summary=result.summary,
        status="SUCCESS",
        file_type="text",
        document_type=request.document_type,
        processing_time_ms=int((time.monotonic() - start_ms) * 1000),
    )

    return SummarizeResponse(
        summary=result.summary,
        steps=result.steps,
        history_id=record.history_id,
    )


@router.post("/summarize/file", response_model=SummarizeResponse)
async def summarize_file_route(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    txt, pdf, docx 파일을 업로드받아 기존 요약 파이프라인으로 처리합니다.

    async def를 쓰는 이유: UploadFile.read()가 비동기 I/O이기 때문입니다.
    """
    # ── 파일 형식 검증 ───────────────────────────────────────────────────────
    filename = file.filename or ""
    ext = ("." + filename.rsplit(".", 1)[-1]).lower() if "." in filename else ""

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail="txt, pdf, docx, png, jpg, jpeg 파일만 업로드할 수 있습니다.",
        )

    pre_steps: list[str] = ["파일 수신 완료"]

    # ── 파일 읽기 및 텍스트 추출 ─────────────────────────────────────────────
    raw = await file.read()

    if ext == ".txt":
        text = _read_txt(raw)
        pre_steps.append("텍스트 추출 완료")
    elif ext == ".pdf":
        try:
            text, pdf_steps = extract_text_from_pdf(raw)
        except RuntimeError as e:
            raise HTTPException(status_code=400, detail=str(e))
        pre_steps.extend(pdf_steps)
    elif ext == ".docx":
        text = _read_docx(raw)
        pre_steps.append("Word 텍스트 추출 완료")
    else:  # .png / .jpg / .jpeg
        try:
            text = extract_text_from_image(raw)
        except RuntimeError as e:
            raise HTTPException(status_code=400, detail=str(e))
        pre_steps.append("이미지 OCR 추출 완료")
        _ocr_preview = text[:300].replace("\n", " ↵ ")
        print(f"[OCR RAW] {_ocr_preview}")
        pre_steps.append(f"[OCR 원문 미리보기] {_ocr_preview}")

    # ── 텍스트 내용 검증 ─────────────────────────────────────────────────────
    if not text:
        raise HTTPException(status_code=400, detail="파일 내용이 비어 있습니다.")

    if len(text) < 10:
        raise HTTPException(
            status_code=400,
            detail="텍스트가 너무 짧습니다. 요약할 내용이 충분한 파일을 업로드해 주세요.",
        )

    # ── 요약 실행 및 이력 저장 ───────────────────────────────────────────────
    summary_mode = "chunked" if len(text) > settings.chunk_threshold else "single"
    file_type = ext.lstrip(".")
    start_ms = time.monotonic()

    try:
        result = summarizer.summarize(text)
    except RuntimeError as e:
        summary_repository.create(
            db,
            user_id=current_user.user_id,
            model_name=settings.ollama_model,
            summary_mode=summary_mode,
            input_chars=len(text),
            output_summary="",
            status="FAILED",
            original_filename=filename,
            file_type=file_type,
            file_size=len(raw),
            error_message=str(e),
            processing_time_ms=int((time.monotonic() - start_ms) * 1000),
        )
        raise HTTPException(status_code=502, detail=str(e))

    record = summary_repository.create(
        db,
        user_id=current_user.user_id,
        model_name=settings.ollama_model,
        summary_mode=summary_mode,
        input_chars=len(text),
        output_summary=result.summary,
        status="SUCCESS",
        original_filename=filename,
        file_type=file_type,
        file_size=len(raw),
        processing_time_ms=int((time.monotonic() - start_ms) * 1000),
    )

    return SummarizeResponse(
        summary=result.summary,
        steps=pre_steps + result.steps,
        history_id=record.history_id,
    )


@router.post("/export")
def export_route(
    request: ExportRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    요약 결과를 지정한 형식의 파일로 변환해 반환합니다.

    요청: { summary, format, source_filename, history_id }
    응답: binary 파일 (Content-Disposition: attachment)

    현재 지원 형식: txt / docx / pdf
    새 형식 추가 방법: export_service._EXPORTERS에 항목 추가 후 exporter 모듈 작성
    """
    try:
        file_bytes, media_type, ext = export_service.export(
            request.summary, request.format, request.source_filename
        )
    except ValueError as e:
        if not request.skip_log:
            download_repository.create(
                db,
                user_id=current_user.user_id,
                download_format=request.format,
                file_name=request.source_filename or None,
                history_id=request.history_id,
                status="FAILED",
                error_message=str(e),
            )
        raise HTTPException(status_code=400, detail=str(e))
    except RuntimeError as e:
        if not request.skip_log:
            download_repository.create(
                db,
                user_id=current_user.user_id,
                download_format=request.format,
                file_name=request.source_filename or None,
                history_id=request.history_id,
                status="FAILED",
                error_message=str(e),
            )
        raise HTTPException(status_code=500, detail=str(e))

    if not request.skip_log:
        download_repository.create(
            db,
            user_id=current_user.user_id,
            download_format=request.format,
            file_name=request.source_filename or None,
            history_id=request.history_id,
            status="SUCCESS",
        )

    filename = _build_export_filename(request.source_filename, ext)

    return Response(
        content=file_bytes,
        media_type=media_type,
        # RFC 5987 방식 — 한글 파일명이 브라우저에서 깨지지 않습니다.
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
    )


@router.post("/chat", response_model=ChatResponse)
def chat_route(
    request: ChatRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    문서 기반 대화 엔드포인트.

    history_id로 요약 기록을 조회하고, 요약문을 컨텍스트로 삼아
    사용자 질문에 답변합니다. 대화 상태는 클라이언트(localStorage)가
    관리하며, 서버는 무상태로 동작합니다.

    Phase 2 확장 시: chat_sessions / chat_messages 테이블을 추가하고
    이 라우트에서 메시지를 DB에 저장하는 로직을 덧붙이면 됩니다.
    서비스 계층(chat_service)은 변경 없이 재사용할 수 있습니다.
    """
    # ── 요약 기록 조회 및 권한 확인 ────────────────────────────────────────────
    record = summary_repository.get_by_id(db, request.history_id)
    if not record:
        raise HTTPException(status_code=404, detail="요약 기록을 찾을 수 없습니다.")
    if record.user_id != current_user.user_id:
        raise HTTPException(status_code=403, detail="접근 권한이 없습니다.")
    if not record.output_summary:
        raise HTTPException(status_code=400, detail="요약 내용이 없는 기록입니다. 요약에 실패한 항목입니다.")

    # ── 질문 검증 ───────────────────────────────────────────────────────────────
    question = request.question.strip()
    if not question:
        raise HTTPException(status_code=400, detail="질문이 비어 있습니다.")
    if len(question) > 1000:
        raise HTTPException(status_code=400, detail="질문이 너무 깁니다. 1,000자 이내로 입력해 주세요.")

    # ── Ollama 호출 ─────────────────────────────────────────────────────────────
    try:
        ai_answer = chat_service.answer(record.output_summary, request.messages, question)
    except RuntimeError as e:
        raise HTTPException(status_code=502, detail=str(e))

    return ChatResponse(answer=ai_answer)
